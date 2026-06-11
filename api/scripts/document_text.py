"""Extract plain text from an uploaded document for the "texto" content source.

The text trigger accepts more than just `.txt`/`.md`: PDFs and Word (`.docx`)
documents are extracted to plain text here so the rest of the pipeline stays
source-agnostic. Dispatch is by file extension first, then magic-byte sniffing
as a fallback for mislabeled or extensionless uploads. The PDF (`pypdf`) and
Word (`python-docx`) parsers are imported lazily so a plain-text upload keeps
working even if those optional packages aren't installed.
"""
from io import BytesIO
from pathlib import Path

# Extensions we route explicitly; anything else falls back to magic-byte sniffing.
_EXT_KIND = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".doc": "doc",
    ".txt": "text",
    ".md": "text",
    ".markdown": "text",
    ".text": "text",
}


def extract_document_text(data: bytes, filename: str = "") -> str:
    """Return the plain-text content of an uploaded document.

    Supports PDF, Word `.docx`, and plain text (`.txt`/`.md`/…). Raises
    `RuntimeError` with a user-facing (Spanish) message for unreadable inputs:
    legacy `.doc`, image-only PDFs, missing optional dependencies, etc.
    """
    if not data:
        return ""
    kind = _detect_kind(data, Path(filename or "").suffix.lower())
    if kind == "pdf":
        return _extract_pdf(data)
    if kind == "docx":
        return _extract_docx(data)
    if kind == "doc":
        raise RuntimeError(
            "El formato .doc (Word antiguo) no es compatible. Abre el documento y "
            "guárdalo como .docx o PDF, y vuelve a subirlo."
        )
    return _decode_text(data)


def _detect_kind(data: bytes, ext: str) -> str:
    """Classify the upload as 'pdf' | 'docx' | 'doc' | 'text'.

    A known extension wins; otherwise sniff the leading magic bytes so a
    mislabeled or extensionless file still routes to the right extractor.
    """
    if ext in _EXT_KIND:
        return _EXT_KIND[ext]
    if data[:5] == b"%PDF-":
        return "pdf"
    if data[:4] == b"PK\x03\x04":  # ZIP container — .docx is the common case
        return "docx"
    if data[:8] == b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1":  # OLE2 compound file — legacy .doc
        return "doc"
    return "text"


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as e:
        raise RuntimeError(
            "Falta la dependencia 'pypdf' para leer PDFs. Instala las dependencias "
            "de la API (pip install -r requirements.txt)."
        ) from e
    try:
        reader = PdfReader(BytesIO(data))
        if reader.is_encrypted:
            # Try the empty password (PDFs encrypted only to restrict editing).
            try:
                reader.decrypt("")
            except Exception:
                raise RuntimeError("El PDF está protegido con contraseña.")
        parts = [(page.extract_text() or "") for page in reader.pages]
    except RuntimeError:
        raise
    except Exception as e:
        raise RuntimeError(f"No se pudo leer el PDF: {e}") from e
    text = "\n".join(p.strip() for p in parts if p.strip())
    if not text.strip():
        raise RuntimeError(
            "El PDF no contiene texto extraíble (puede ser un PDF escaneado o solo "
            "imágenes). Usa un PDF con texto seleccionable."
        )
    return text


def _extract_docx(data: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as e:
        raise RuntimeError(
            "Falta la dependencia 'python-docx' para leer archivos Word. Instala las "
            "dependencias de la API (pip install -r requirements.txt)."
        ) from e
    try:
        document = docx.Document(BytesIO(data))
    except Exception as e:
        raise RuntimeError(f"No se pudo leer el archivo Word (.docx): {e}") from e
    parts = [p.text for p in document.paragraphs]
    # Tables aren't part of `paragraphs`; include their cell text so structured
    # docs (briefs, outlines) don't lose content.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    text = "\n".join(p.strip() for p in parts if p and p.strip())
    if not text.strip():
        raise RuntimeError("El documento Word está vacío o no contiene texto.")
    return text


def _decode_text(data: bytes) -> str:
    """Decode raw bytes as text, trying the common encodings before giving up.

    `utf-8-sig` strips a BOM if present; `cp1252`/`latin-1` cover text exported
    from Windows editors. The final pass replaces undecodable bytes so we never
    raise on a near-text file.
    """
    for enc in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")
